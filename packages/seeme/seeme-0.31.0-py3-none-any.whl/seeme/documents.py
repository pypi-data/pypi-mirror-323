import os
import os.path
import sys
import sched, time
import shutil
import glob
import time
from datetime import datetime
import uuid

from datetime import datetime

from typing import Union

import docx
from PyPDF2 import PdfReader

def format_pdf_date(date_str):
    """
    Convert PDF date string to readable format.
    PDF dates are typically in format "D:YYYYMMDDHHmmSS" or "D:YYYYMMDDHHmmSS+HH'mm'"
    
    Args:
        date_str (str): PDF date string
        
    Returns:
        str: Formatted date string or None if conversion fails
    """
    if not date_str:
        return None
        
    try:
        # Remove 'D:' prefix if present
        if date_str.startswith('D:'):
            date_str = date_str[2:]
        
        # Basic format: YYYYMMDDHHmmSS
        if len(date_str) >= 14:
            dt = datetime.strptime(date_str[:14], '%Y%m%d%H%M%S')
            return dt.strftime('%Y-%m-%d %H:%M:%S')
        return None
    except:
        return None


def get_word_text(filename:str, metadata:dict) -> list[dict]:
    """
    Get all text from a Word file, using the package 'docx'. No images are currently OCR'd.
    
    Args:
        filename (str): the full filename to use

        metadata (dict): the metadata for the file, currently we only use the 'id' value.
    
    Returns:
        list[dict]: A list of dictionary, where every dictionary is the text for one of the paragraphs
    """
    doc = docx.Document(filename)
    jsons = []
    for para in doc.paragraphs:
        try:
            line = {
                "text": para.text,
                "id": metadata["id"],
            }
            
            jsons.append(line)
        except:
            print(f"File could not be processed: {filename}")

    return jsons


def get_pdf_text(filename:str, metadata:dict) -> list[dict]:
    """
    Get all text from a PDF file, using the package 'pymupdf'. No images or image PDFs are currently OCR'd.
    
    Args:
        filename (str): the full filename to use

        metadata (dict): the metadata for the file, currently we only use the 'id' value.
    
    Returns:
        list[dict]: A list of dictionary, where every dictionary is the text for one of the pages
    """
    #doc = fitz.open(filename)
    reader = PdfReader(filename)
    jsons = []
    #for page in doc:
    for page in reader.pages:
        try:
            line = {
                #"text": page.get_text(),
                "text": page.extract_text(),
                "id": metadata["id"]
            }
            
            jsons.append(line)
        except:
            print(f"File could not be processed: {filename}")

    return jsons

def get_filename_metadata(filename:str, metadata_id:[str, int]):
    """
    Get the filename metadata.
    
    Args:
        filename (str):    The input filename to extract the metadata from. 
                 
        metadata_id ([str, int]): The 'id' to be used for the metadata
    
    Returns:
        metadata = {
            "id": metadata_id,
            "filename": "the_actual_filename"
        }
    """


    actual_filename= os.path.basename(filename)
    metadata = {
        "id": metadata_id,
        "filename": actual_filename
    }
        
    return metadata

def process_filename(filename:str, metadata_id:Union[str, int]) -> tuple[list[dict], dict]:
    """
    Get all texts and metadata from a PDF/Word file.
    
    Args:
        file (str): the full filename to use

        metadata_id ([str, int]): the metadata_id to be used
    
    Returns:
        - list[dict]: A list of dicts with text and and the metadata id. 
        - dict: The created metadata
    """
    try:
        metadata = get_filename_metadata(filename, metadata_id)
       
        read_inputs = ""
        
        if filename.endswith(".docx") or filename.endswith(".doc"):
            #print("word doc")
            doc = docx.Document(filename)
            core_properties = doc.core_properties
            metadata['created'] = str(core_properties.created)
            metadata['modified'] = str(core_properties.modified)
            metadata['author'] = core_properties.author
            read_inputs = get_word_text(filename, metadata)
       
        elif filename.endswith(".pdf") :
            #print("pdf doc")
            #doc = fitz.open(filename)
            reader = PdfReader(filename)
            pdf_metadata = reader.metadata
            metadata['created'] =  str(format_pdf_date(pdf_metadata.get('/CreationDate')))
            metadata['modified'] = str(format_pdf_date(pdf_metadata.get('/ModDate', None)))
            metadata['author'] = pdf_metadata.author
            read_inputs = get_pdf_text(filename, metadata)
    except Exception as e:
        print(f"Error handling file: {filename}")
        print(e)
        return "", metadata
        
    return read_inputs, metadata